import pandas as pd
from docx import Document


def parse_description_file(file_path: str) -> pd.DataFrame:
    """
    Parse the description .docx file and return a DataFrame.
    """
    doc = Document(file_path)

    if not doc.tables:
        raise ValueError("No tables found in the description document.")

    # Find the first table that contains 2-column key-value rows
    table = None
    for candidate in doc.tables:
        if any(len(row.cells) >= 2 and row.cells[0].text.strip() for row in candidate.rows):
            table = candidate
            break

    if table is None:
        raise ValueError(
            f"No suitable key-value table found. "
            f"Document has {len(doc.tables)} table(s) but none with 2-column rows."
        )

    data = {}
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        key = row.cells[0].text.strip()
        value = row.cells[1].text.strip()
        # Skip merged/title rows (e.g. "English form" header) where both cells hold the same text
        if key == value:
            continue
        if key.lower() == "property features":
            value = "; ".join(
                line.strip().lstrip("•·●■-– ").strip() for line in value.splitlines() if line.strip()
            )
        data[key] = value

    if not data:
        row_summary = [(len(r.cells), [c.text.strip()[:40] for c in r.cells]) for r in table.rows]
        raise ValueError(
            f"No key-value data found in document table. "
            f"Table has {len(table.rows)} row(s), {len(table.columns)} column(s). "
            f"Row structure: {row_summary}"
        )

    publishing_properties = pd.DataFrame(list(data.items()), columns=["Key", "Value"])
    publishing_properties = publishing_properties.set_index("Key").T.reset_index(drop=True)
    publishing_properties.columns = publishing_properties.columns.map(str.lower)

    print(f"Parsed properties:\n{publishing_properties.to_string()}")

    return publishing_properties
